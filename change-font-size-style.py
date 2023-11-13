
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def read_word_file(input_path):
    doc = Document(input_path)
    text = ""

    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    return text


def modify_word_file(input_path, output_path, new_font_size, new_font_style):
    # Load the existing Word document
    doc = Document(input_path)
    

    # Function to modify font size and style for runs in a paragraph
    
    def modify_runs(paragraph):
        for run in paragraph.runs:
            # Change font size
            run.font.size = Pt(new_font_size)
            # Change font style
            run.font.name = new_font_style

    # Iterate through paragraphs and modify font size and style
    for paragraph in doc.paragraphs:
        modify_runs(paragraph)

    # Iterate through tables and modify font size and style for text in cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    modify_runs(paragraph)

    # Save the modified document to a new file
    doc.save(output_path)

if __name__ == "__main__":
    # Specify the input and output file paths
    input_file ="Enrollment Challenges Recruiting.docx"
    output_file = "Enrollment Challenges Recruiting - Converted.docx"

    # Specify the new font size and font style
    font_size = 9  # The desired font size
    font_style = 'Bodoni MT Black'

    # Call the function to modify the Word file
    modify_word_file(input_file, output_file, font_size, font_style)

    document_text = read_word_file(input_file)

# Print the extracted text
    print(document_text)
    print(f"Word file modified and saved to {output_file}")
